import re
import io
from typing import List, Dict, Any

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXParser:
    SUPPORTED_EXTENSIONS = [".docx", ".doc", ".docm", ".xlsm", ".pptm"]

    MACRO_EXTENSIONS = {".docm", ".xlsm", ".pptm"}

    @staticmethod
    def parse(file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
        urls = []
        text = ""
        metadata = {}
        scripts = []
        macros_detected = False

        if not DOCX_AVAILABLE:
            text = DOCXParser._extract_text_raw(file_bytes)
            urls = DOCXParser._extract_urls_from_text(text)
            macros_detected = DOCXParser._check_macro_indicators(file_bytes)
            return {
                "text": text,
                "urls": urls,
                "metadata": metadata,
                "scripts": scripts,
                "macros_detected": macros_detected,
            }

        ext = "." + filename.split(".")[-1].lower() if filename else ""

        try:
            with io.BytesIO(file_bytes) as f:
                doc = Document(f)

                text = "\n".join([para.text for para in doc.paragraphs])

                for para in doc.paragraphs:
                    for run in para.runs:
                        if run.hyperlink:
                            urls.append(run.hyperlink.address or "")

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"

                urls.extend(DOCXParser._extract_urls_from_text(text))
                urls = [u for u in urls if u]
                urls = list(set(urls))

                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                }

                if hasattr(doc, "core_properties"):
                    props = doc.core_properties
                    metadata.update(
                        {
                            "author": props.author or "",
                            "title": props.title or "",
                            "subject": props.subject or "",
                            "created": str(props.created) if props.created else "",
                            "modified": str(props.modified) if props.modified else "",
                        }
                    )

                macros_detected = DOCXParser._check_macro_indicators(file_bytes)

                if macros_detected:
                    scripts.append("VBA macro detected in document")

        except Exception as e:
            text = DOCXParser._extract_text_raw(file_bytes)
            urls = DOCXParser._extract_urls_from_text(text)
            macros_detected = DOCXParser._check_macro_indicators(file_bytes)

        return {
            "text": text,
            "urls": urls,
            "metadata": metadata,
            "scripts": scripts,
            "macros_detected": macros_detected,
        }

    @staticmethod
    def _extract_text_raw(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except:
            return ""

    @staticmethod
    def _extract_urls_from_text(text: str) -> List[str]:
        url_pattern = re.compile(
            r"https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s<>()\[\]]*", re.IGNORECASE
        )
        return url_pattern.findall(text)

    @staticmethod
    def _check_macro_indicators(file_bytes: bytes) -> bool:
        macro_indicators = [
            b"vbaProject.bin",
            b"_rels/vbaProject.bin.rels",
            b"word/vbaProject.bin",
            b"ppt/vbaProject.bin",
            b"xl/vbaProject.bin",
            b"MSUBST",
            b"chrpdevice",
            b"_VBA_PROJECT",
        ]

        for indicator in macro_indicators:
            if indicator in file_bytes:
                return True

        return False

    @staticmethod
    def _detect_suspicious_content(text: str) -> List[str]:
        scripts = []
        suspicious_patterns = [
            r"powershell\s+-",
            r"cmd\.exe",
            r"certutil",
            r"bitsadmin",
            r"wscript",
            r"cscript",
            r"reg\s+add",
            r"schtasks",
            r"base64",
            r"encodedcommand",
            r"downloadfile",
            r"invoke-webrequest",
            r"curl\s+",
            r"wget\s+",
        ]

        for pattern in suspicious_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                scripts.append(pattern)

        return scripts
